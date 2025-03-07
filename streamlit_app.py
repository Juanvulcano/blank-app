import zipfile
from io import BytesIO

import docx
import openpyxl
import pandas as pd
import streamlit as st
from google.cloud import translate_v2 as translate
from google.oauth2 import service_account

# Google Cloud Translation API setup
credentials = service_account.Credentials.from_service_account_info(
    st.secrets["gcp_service_account"]
)
translate_client = translate.Client(credentials=credentials)

# Function to translate text with exclusions


def translate_text(text, target_language, model="Google", exclude_words=None):
    if exclude_words:
        # Sort the words by length to avoid partial replacements
        exclude_words = sorted(exclude_words, key=len, reverse=True)
        word_positions = {}

        # Find and mark the positions of excluded words
        for word in exclude_words:
            word = word.strip()
            if word in text:
                start_index = text.find(word)
                end_index = start_index + len(word)
                word_positions[word] = (start_index, end_index)
                text = text.replace(word, f"<EXCLUDE_{word}>")  # Mark the word

    # Translate the text
    if model == "Google":
        result = translate_client.translate(
            text, target_language=target_language
        )
        translated_text = result["translatedText"]
    else:
        translated_text = "Unsupported translation model."

    # Replace markers with the original words
    if exclude_words:
        for word in word_positions:
            translated_text = translated_text.replace(
                f"<EXCLUDE_{word}>", word)

    return translated_text

# Function to handle file upload and translation


def handle_file_upload(file, target_languages, model="Google", exclude_words=None):
    translations = {}

    for language in target_languages:
        if file.type == "text/plain":
            content = file.read().decode("utf-8")
            translated_content = translate_text(
                content, languages[language], model, exclude_words)
            translations[language] = translated_content.encode('utf-8')
        elif file.type == "text/csv":
            df = pd.read_csv(file)
            translated_df = df.applymap(
                lambda x: translate_text(
                    x, languages[language], model, exclude_words)
            )
            translations[language] = translated_df.to_csv(
                index=False).encode('utf-8')
        elif file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(file)
            translated_df = df.applymap(
                lambda x: translate_text(
                    x, languages[language], model, exclude_words)
            )
            output = BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                translated_df.to_excel(
                    writer, index=False, sheet_name="Sheet1")
                writer.save()
            translations[language] = output.getvalue()
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)

            # Translate paragraphs
            for para in doc.paragraphs:
                para.text = translate_text(
                    para.text, languages[language], model, exclude_words)

            # Translate tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = translate_text(
                            cell.text, languages[language], model, exclude_words)

            output = BytesIO()
            doc.save(output)
            translations[language] = output.getvalue()
        else:
            translations[language] = "Unsupported file type.".encode('utf-8')

        # Reset file pointer to the beginning for the next translation
        file.seek(0)

    return translations

# Function to create a ZIP file


def create_zip(translations, original_filename):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file:
        for language, content in translations.items():
            filename = f"{language}_{original_filename}"
            zip_file.writestr(filename, content)
    zip_buffer.seek(0)
    return zip_buffer


# Streamlit app UI
st.title("Document Translator")

uploaded_file = st.file_uploader(
    "Upload a document", type=["txt", "csv", "xlsx", "docx"]
)

languages = {
    "English (US)": "en",
    "English (UK)": "en-GB",
    "English (AU)": "en-AU",
    "Spanish (ES)": "es",
    "Spanish (MX)": "es-MX",
    "French": "fr",
    "German (DE)": "de",
    "Japanese": "ja",
    "Portuguese (BR)": "pt-BR",
    "Portuguese (PT)": "pt-PT",
    "Italian": "it",
    "Turkish": "tr",
    "Polish": "pl",
    "Dutch (NL)": "nl",
    "Chinese (Simplified)": "zh-CN",
    "Korean": "ko",
    "Russian": "ru",
    "Arabic": "ar",
}

models = ["Google", "DeepL"]  # Add more models here as needed

selected_languages = st.multiselect(
    "Select target languages", list(languages.keys()))
model = st.selectbox("Select translation model", models)

exclude_words_input = st.text_area(
    "Enter words/phrases to exclude from translation (comma-separated):",
    value="Apple Watch, apple watch")

if st.button("Translate"):
    if uploaded_file is not None:
        exclude_words = [word.strip() for word in exclude_words_input.split(
            ',')] if exclude_words_input else None
        translations = handle_file_upload(
            uploaded_file, selected_languages, model, exclude_words)

        if translations:
            zip_file = create_zip(translations, uploaded_file.name)
            st.download_button(
                label="Download Translated Documents",
                data=zip_file,
                file_name="translated_documents.zip",
                mime="application/zip"
            )
    else:
        st.error("Please upload a file first.")
